"""
Word Document Report Generation for JAXSR DOE Studies.

Generates formatted ``.docx`` reports containing study summaries,
model results, coefficient tables, and embedded plots.

Requires the ``reports`` optional dependency group::

    pip install jaxsr[reports]

Which installs ``python-docx``.
"""

from __future__ import annotations

import io
from typing import TYPE_CHECKING

import numpy as np

if TYPE_CHECKING:
    from .study import DOEStudy


def _check_docx():
    """Raise ImportError with install instructions if python-docx is missing."""
    try:
        import docx  # noqa: F401
    except ImportError:
        raise ImportError(
            "python-docx is required for Word report generation. "
            "Install it with: pip install jaxsr[reports]"
        ) from None


def generate_word_report(study: DOEStudy, filepath: str) -> str:
    """
    Generate a Word document report for a fitted DOE study.

    The report includes:

    - Title page with study name, description, and metadata
    - Factor summary table
    - Model expression and metrics
    - Coefficient table
    - Prediction vs actual table
    - Embedded residual and prediction plots (matplotlib)
    - Iteration history (if available)

    Parameters
    ----------
    study : DOEStudy
        A fitted study with observations.
    filepath : str
        Path for the output ``.docx`` file.

    Returns
    -------
    filepath : str
        The path that was written (same as input).

    Raises
    ------
    RuntimeError
        If the study has no fitted model or no observations.
    ImportError
        If python-docx is not installed.
    """
    _check_docx()
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Inches, Pt, RGBColor

    if not study.is_fitted:
        raise RuntimeError("Study has no fitted model. Call study.fit() first.")
    if study.X is None or study.y is None:
        raise RuntimeError("Study has no observations.")

    model = study.model
    result = model._result

    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # =========================================================================
    # Title
    # =========================================================================
    title = doc.add_heading(f"DOE Study Report: {study.name}", level=0)
    title.runs[0].font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    if study.description:
        doc.add_paragraph(study.description)

    doc.add_paragraph(
        f"Created: {study.meta.get('created', 'unknown')}  |  "
        f"Modified: {study.meta.get('modified', 'unknown')}  |  "
        f"JAXSR v{study.meta.get('jaxsr_version', 'unknown')}"
    ).runs[0].font.size = Pt(9)

    # =========================================================================
    # Study Overview
    # =========================================================================
    doc.add_heading("Study Overview", level=1)

    table = doc.add_table(rows=1, cols=4, style="Light Shading Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Factor", "Type", "Bounds / Levels", ""]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for i, name in enumerate(study.factor_names):
        row = table.add_row().cells
        row[0].text = name

        ftype = "continuous"
        if study.feature_types and i < len(study.feature_types):
            ftype = study.feature_types[i]
        row[1].text = ftype

        if ftype == "categorical" and study.categories and i in study.categories:
            row[2].text = ", ".join(str(lev) for lev in study.categories[i])
        else:
            low, high = study.bounds[i]
            row[2].text = f"[{low}, {high}]"

    p = doc.add_paragraph()
    p.add_run("Observations: ").bold = True
    p.add_run(str(study.n_observations))

    if study._design_config:
        p = doc.add_paragraph()
        p.add_run("Design method: ").bold = True
        p.add_run(study._design_config.get("method", "unknown"))

    # =========================================================================
    # Model Results
    # =========================================================================
    doc.add_heading("Model Results", level=1)

    doc.add_heading("Expression", level=2)
    expr_para = doc.add_paragraph()
    expr_run = expr_para.add_run(f"y = {result.expression()}")
    expr_run.font.size = Pt(12)
    expr_run.bold = True

    doc.add_heading("Model Metrics", level=2)

    # Compute R-squared
    predict_fn = model.to_callable()
    X_np = np.asarray(study.X)
    y_actual = np.asarray(study.y)
    y_pred = predict_fn(X_np)
    residuals = y_actual - y_pred
    ss_res = float(np.sum(residuals**2))
    ss_tot = float(np.sum((y_actual - np.mean(y_actual)) ** 2))
    r_squared = 1.0 - ss_res / ss_tot if ss_tot > 0 else 0.0

    metrics_table = doc.add_table(rows=1, cols=2, style="Light Shading Accent 1")
    metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(metrics_table.rows[0].cells, ["Metric", "Value"], strict=True):
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    metrics = [
        ("R\u00b2", f"{r_squared:.6f}"),
        ("MSE", f"{result.mse:.6g}"),
        ("RMSE", f"{np.sqrt(result.mse):.6g}"),
        ("AIC", f"{result.aic:.4f}"),
        ("BIC", f"{result.bic:.4f}"),
        ("AICc", f"{result.aicc:.4f}"),
        ("Number of Terms", str(result.n_terms)),
        ("Observations", str(result.n_samples)),
    ]
    for label, value in metrics:
        row = metrics_table.add_row().cells
        row[0].text = label
        row[1].text = value

    # =========================================================================
    # Coefficients Table
    # =========================================================================
    doc.add_heading("Coefficients", level=2)

    coeff_table = doc.add_table(rows=1, cols=3, style="Light Shading Accent 1")
    coeff_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(
        coeff_table.rows[0].cells,
        ["Basis Function", "Coefficient", "|Coefficient|"],
        strict=True,
    ):
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    coeffs = np.array(result.coefficients)
    for name, coeff in zip(result.selected_names, coeffs, strict=True):
        row = coeff_table.add_row().cells
        row[0].text = name
        row[1].text = f"{float(coeff):.6g}"
        row[2].text = f"{float(abs(coeff)):.6g}"

    # =========================================================================
    # Plots
    # =========================================================================
    doc.add_heading("Diagnostic Plots", level=1)

    try:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        # Predicted vs Actual
        fig, ax = plt.subplots(1, 1, figsize=(5, 4.5))
        ax.scatter(y_actual, y_pred, alpha=0.7, edgecolors="k", linewidths=0.5)
        min_val = min(y_actual.min(), y_pred.min())
        max_val = max(y_actual.max(), y_pred.max())
        ax.plot([min_val, max_val], [min_val, max_val], "r--", linewidth=1)
        ax.set_xlabel("Actual")
        ax.set_ylabel("Predicted")
        ax.set_title(f"Predicted vs Actual (R\u00b2 = {r_squared:.4f})")
        ax.grid(True, alpha=0.3)
        fig.tight_layout()

        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150)
        plt.close(fig)
        buf.seek(0)

        doc.add_picture(buf, width=Inches(4.5))
        buf.close()

        # Residuals vs Predicted
        fig, ax = plt.subplots(1, 1, figsize=(5, 4.5))
        ax.scatter(y_pred, residuals, alpha=0.7, edgecolors="k", linewidths=0.5)
        ax.axhline(y=0, color="r", linestyle="--", linewidth=1)
        ax.set_xlabel("Predicted")
        ax.set_ylabel("Residual")
        ax.set_title("Residuals vs Predicted")
        ax.grid(True, alpha=0.3)
        fig.tight_layout()

        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150)
        plt.close(fig)
        buf.seek(0)

        doc.add_picture(buf, width=Inches(4.5))
        buf.close()

        # Coefficient magnitudes bar chart
        fig, ax = plt.subplots(1, 1, figsize=(6, max(3, len(result.selected_names) * 0.4)))
        sorted_indices = np.argsort(np.abs(coeffs))
        sorted_names = [result.selected_names[i] for i in sorted_indices]
        sorted_abs = np.abs(coeffs[sorted_indices])
        ax.barh(range(len(sorted_names)), sorted_abs, color="#4472C4")
        ax.set_yticks(range(len(sorted_names)))
        ax.set_yticklabels(sorted_names)
        ax.set_xlabel("|Coefficient|")
        ax.set_title("Coefficient Magnitudes")
        ax.grid(True, alpha=0.3, axis="x")
        fig.tight_layout()

        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150)
        plt.close(fig)
        buf.seek(0)

        doc.add_picture(buf, width=Inches(5))
        buf.close()

    except ImportError:
        doc.add_paragraph(
            "Matplotlib is not available. Plots could not be generated. "
            "Install matplotlib for plot support."
        )

    # =========================================================================
    # Prediction Table
    # =========================================================================
    doc.add_heading("Predictions", level=1)

    pred_table = doc.add_table(rows=1, cols=4, style="Light Shading Accent 1")
    pred_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(
        pred_table.rows[0].cells,
        ["Run", "Actual", "Predicted", "Residual"],
        strict=True,
    ):
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for i in range(len(y_actual)):
        row = pred_table.add_row().cells
        row[0].text = str(i + 1)
        row[1].text = f"{float(y_actual[i]):.4f}"
        row[2].text = f"{float(y_pred[i]):.4f}"
        row[3].text = f"{float(residuals[i]):.4f}"

    # =========================================================================
    # Iteration History
    # =========================================================================
    if study.iterations:
        doc.add_heading("Iteration History", level=1)

        hist_table = doc.add_table(rows=1, cols=4, style="Light Shading Accent 1")
        hist_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for cell, text in zip(
            hist_table.rows[0].cells,
            ["Round", "Points Added", "Model", "Notes"],
            strict=True,
        ):
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for it in study.iterations:
            row = hist_table.add_row().cells
            row[0].text = str(it.round_number)
            row[1].text = str(it.n_points_added)
            row[2].text = it.model_expression or "\u2014"
            row[3].text = it.notes or ""

    doc.save(filepath)
    return filepath
