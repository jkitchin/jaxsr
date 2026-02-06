"""Tests for Word document report generation."""

from __future__ import annotations

import numpy as np
import pytest

from jaxsr.reporting import generate_word_report
from jaxsr.study import DOEStudy


@pytest.fixture
def fitted_study():
    """A fitted study for report testing."""
    np.random.seed(42)
    study = DOEStudy(
        name="report_test",
        factor_names=["x1", "x2"],
        bounds=[(0, 1), (0, 1)],
        description="Testing Word report generation",
    )
    X = np.random.rand(30, 2)
    y = 2.0 * X[:, 0] + 3.0 * X[:, 1] + 0.01 * np.random.randn(30)
    study.add_observations(X, y, notes="Initial data")
    study.fit(max_terms=3)
    return study


@pytest.fixture
def fitted_study_with_iterations():
    """A fitted study with multiple iterations."""
    np.random.seed(42)
    study = DOEStudy(
        name="iteration_report",
        factor_names=["x1", "x2"],
        bounds=[(0, 1), (0, 1)],
    )
    # First batch
    X1 = np.random.rand(15, 2)
    y1 = 2.0 * X1[:, 0] + 3.0 * X1[:, 1] + 0.01 * np.random.randn(15)
    study.add_observations(X1, y1, notes="Batch 1")
    study.fit(max_terms=3)

    # Second batch
    X2 = np.random.rand(15, 2)
    y2 = 2.0 * X2[:, 0] + 3.0 * X2[:, 1] + 0.01 * np.random.randn(15)
    study.add_observations(X2, y2, notes="Batch 2")
    study.fit(max_terms=3)

    return study


class TestGenerateWordReport:
    """Test Word document report generation."""

    def test_generates_file(self, fitted_study, tmp_path):
        filepath = str(tmp_path / "report.docx")
        result = generate_word_report(fitted_study, filepath)
        assert result == filepath

        from docx import Document

        doc = Document(filepath)
        assert len(doc.paragraphs) > 0

    def test_report_contains_title(self, fitted_study, tmp_path):
        filepath = str(tmp_path / "report.docx")
        generate_word_report(fitted_study, filepath)

        from docx import Document

        doc = Document(filepath)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "report_test" in text

    def test_report_contains_expression(self, fitted_study, tmp_path):
        filepath = str(tmp_path / "report.docx")
        generate_word_report(fitted_study, filepath)

        from docx import Document

        doc = Document(filepath)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "y =" in text

    def test_report_has_tables(self, fitted_study, tmp_path):
        filepath = str(tmp_path / "report.docx")
        generate_word_report(fitted_study, filepath)

        from docx import Document

        doc = Document(filepath)
        # Should have: factor table, metrics table, coefficients table, predictions table
        assert len(doc.tables) >= 4

    def test_report_has_images(self, fitted_study, tmp_path):
        """Report should contain embedded matplotlib plots."""
        filepath = str(tmp_path / "report.docx")
        generate_word_report(fitted_study, filepath)

        from docx import Document

        doc = Document(filepath)
        # Check for inline shapes (images)
        image_count = sum(
            1
            for p in doc.paragraphs
            for r in p.runs
            if r._element.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
            )
        )
        # At least the predicted vs actual and residual plots
        assert image_count >= 2

    def test_report_with_iterations(self, fitted_study_with_iterations, tmp_path):
        filepath = str(tmp_path / "report.docx")
        generate_word_report(fitted_study_with_iterations, filepath)

        from docx import Document

        doc = Document(filepath)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Iteration" in text

    def test_not_fitted_raises(self, tmp_path):
        study = DOEStudy(name="empty", factor_names=["x"], bounds=[(0, 1)])
        with pytest.raises(RuntimeError, match="no fitted model"):
            generate_word_report(study, str(tmp_path / "bad.docx"))

    def test_report_contains_description(self, fitted_study, tmp_path):
        filepath = str(tmp_path / "report.docx")
        generate_word_report(fitted_study, filepath)

        from docx import Document

        doc = Document(filepath)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Testing Word report generation" in text

    def test_report_contains_metrics(self, fitted_study, tmp_path):
        filepath = str(tmp_path / "report.docx")
        generate_word_report(fitted_study, filepath)

        from docx import Document

        doc = Document(filepath)
        # Find the metrics table
        found_mse = False
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                if "MSE" in cells:
                    found_mse = True
                    break
        assert found_mse
